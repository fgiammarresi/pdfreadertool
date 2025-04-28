import subprocess
import time
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTImage, LTTextBoxHorizontal
from docx import Document
from docx.shared import Inches
import os  # Importa il modulo os

def leggi_pdf(percorso_file):
    """
    Legge il contenuto di un file PDF, estraendo testo, tabelle (migliorate) e gestendo le immagini.
    Restituisce una lista di elementi estratti.
    """
    elementi = []
    try:
        for page_layout in extract_pages(percorso_file):
            righe_tabella = {}  # Dizionario per tracciare il testo nelle righe
            for element in page_layout:
                if isinstance(element, LTTextContainer) and not isinstance(element, LTTextBoxHorizontal):
                    elementi.append({'type': 'text', 'content': element.get_text()})
                elif isinstance(element, LTTextBoxHorizontal):
                    y0 = round(element.bbox[1], 2)  # Arrotonda la coordinata y per raggruppare il testo nella stessa riga
                    if y0 not in righe_tabella:
                        righe_tabella[y0] = []
                    righe_tabella[y0].append(element)
                elif isinstance(element, LTImage):
                    elementi.append({'type': 'image', 'content': element})

            # Ordina il testo in ogni riga per coordinata x
            for y in righe_tabella:
                righe_tabella[y] = sorted(righe_tabella[y], key=lambda x: x.bbox[0])

            # Costruisci le righe della tabella (semplice esempio)
            for y in sorted(righe_tabella.keys(), reverse=True):  # Ordina le righe dalla cima alla fine della pagina
                riga_testo = " ".join(elemento.get_text().strip() for elemento in righe_tabella[y])
                elementi.append({'type': 'table_row', 'content': riga_testo})  # Aggiungi come riga di tabella

    except FileNotFoundError:
        print(f"Errore: File non trovato nel percorso '{percorso_file}'")
        return None
    return elementi

def crea_docx(elementi, nome_file_output="output.docx"):
    """
    Crea un file DOCX dagli elementi estratti dal PDF.
    """
    documento = Document()
    for elemento in elementi:
        if elemento['type'] == 'text':
            documento.add_paragraph(elemento['content'])
        elif elemento['type'] == 'table_row':
            documento.add_paragraph(elemento['content'])  # Aggiungi semplicemente la riga di testo
        elif elemento['type'] == 'image':
            # Per ora, proviamo ad aggiungere l'immagine se abbiamo il nome del file
            if hasattr(elemento['content'], 'stream') and hasattr(elemento['content'], 'name'):
                try:
                    with open(elemento['content'].name, 'wb') as f:
                        f.write(elemento['content'].stream.read())
                    documento.add_picture(elemento['content'].name, width=Inches(4.0))  # Regola la larghezza a piacimento
                    os.remove(elemento['content'].name)  # Pulisci il file temporaneo
                except Exception as e:
                    print(f"Errore nell'aggiungere l'immagine: {e}")
            else:
                documento.add_paragraph("Immagine non visualizzata.")

    documento.save(nome_file_output)
    print(f"File DOCX '{nome_file_output}' creato con successo.")
    # files.download(nome_file_output) # Funziona solo in Colab (rimosso)

    # Aggiungiamo un controllo per assicurarci che nome_file_output sia valido.
    # Questo è un controllo di sicurezza per prevenire errori successivi.
    if nome_file_output:
        # Se nome_file_output ha un valore (cioè, non è vuoto o None), lo restituiamo.
        return nome_file_output
    else:
        # Altrimenti, se nome_file_output è vuoto o None, stampiamo un messaggio di errore
        # per segnalare il problema e restituiamo None.
        print("Errore: Il nome del file di output è vuoto.")
        return None

def main():
    percorso_pdf = input("Inserisci il percorso del file PDF da elaborare: ")

    print("\nSeleziona l'attività:")
    print("1. Trascrizione")
    print("2. Traduzione (TBD)")
    print("3. Sintesi (TBD)")
    print("4. Nuove funzionalità TBD")
    scelta_attivita = input("Inserisci il numero dell'attività desiderata: ")

    if scelta_attivita == '1':
        print("\nSeleziona il formato di output:")
        print("1. DOCX")
        scelta_formato = input("Inserisci il numero del formato desiderato: ")
        if scelta_formato == '1':
            print("\nInizio processo di trascrizione...")
            start_time = time.time()
            elementi_estratti = leggi_pdf(percorso_pdf)
            if elementi_estratti:
                crea_docx(elementi_estratti)
            end_time = time.time()
            elapsed_time = end_time - start_time
            print(f"Tempo impiegato: {elapsed_time:.2f} secondi.")
        else:
            print("Formato di output non supportato per ora.")
    else:
        print("Attività non implementata.")

if __name__ == "__main__":
    main()
